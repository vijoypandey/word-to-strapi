#!/usr/bin/env python3
"""
Word to Strapi Converter
Converts Word documents following a blog template to Strapi layouts
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import json
from docx import Document
import re
import os
from datetime import datetime


class WordToStrapiConverter:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Word to Strapi Converter")
        self.root.geometry("900x700")
        
        # Modern color scheme
        self.colors = {
            'bg_primary': '#1a1a2e',      # Dark blue-gray
            'bg_secondary': '#16213e',     # Slightly lighter blue-gray
            'bg_accent': '#0f3460',        # Dark blue accent
            'text_primary': '#ffffff',      # White text
            'text_secondary': '#e8e8e8',   # Light gray text
            'accent_primary': '#00d4aa',   # Teal accent
            'accent_secondary': '#ff6b6b', # Coral accent
            'button_primary': '#00d4aa',   # Teal button
            'button_secondary': '#ff6b6b', # Coral button
            'success': '#4CAF50',          # Green
            'warning': '#ff9800',          # Orange
            'error': '#f44336'             # Red
        }
        
        self.root.configure(bg=self.colors['bg_primary'])
        
        # Template fields
        self.template_fields = [
            "Working Title",
            "Author", 
            "Topic",
            "Blog Category",
            "Target Keywords",
            "Target Audience",
            "Funnel Stage",
            "CTA",
            "Working Meta Description"
        ]
        
        self.setup_ui()
        self.setup_hover_effects()
    
    def setup_hover_effects(self):
        """Setup hover effects for buttons"""
        def on_enter(event):
            event.widget.configure(bg=self.colors['accent_primary'])
            
        def on_leave(event):
            if event.widget == self.browse_btn:
                event.widget.configure(bg=self.colors['button_primary'])
            elif event.widget == self.convert_btn:
                event.widget.configure(bg=self.colors['button_secondary'])
        
        # Bind hover effects
        self.browse_btn.bind("<Enter>", on_enter)
        self.browse_btn.bind("<Leave>", on_leave)
        self.convert_btn.bind("<Enter>", on_enter)
        self.convert_btn.bind("<Leave>", on_leave)
    
    def setup_ui(self):
        """Setup the user interface"""
        # Main frame
        main_frame = tk.Frame(self.root, bg=self.colors['bg_primary'])
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Title with gradient effect
        title_frame = tk.Frame(main_frame, bg=self.colors['bg_primary'])
        title_frame.pack(fill=tk.X, pady=(0, 20))
        
        title_label = tk.Label(title_frame, text="Word to Strapi Converter", 
                             font=("Helvetica", 24, "bold"), 
                             bg=self.colors['bg_primary'], 
                             fg=self.colors['accent_primary'])
        title_label.pack()
        
        subtitle_label = tk.Label(title_frame, text="Convert Word documents to Strapi layouts", 
                                font=("Helvetica", 12), 
                                bg=self.colors['bg_primary'], 
                                fg=self.colors['text_secondary'])
        subtitle_label.pack()
        
        # File selection frame
        file_frame = tk.LabelFrame(main_frame, text="Select Word Document", 
                                 bg=self.colors['bg_secondary'], 
                                 fg=self.colors['text_primary'],
                                 font=("Helvetica", 12, "bold"))
        file_frame.pack(fill=tk.X, pady=(0, 20))
        
        # File path display
        self.file_path_var = tk.StringVar()
        self.file_path_entry = tk.Entry(file_frame, textvariable=self.file_path_var, 
                                      width=60, font=("Helvetica", 10),
                                      bg=self.colors['bg_accent'], 
                                      fg=self.colors['text_primary'],
                                      insertbackground=self.colors['accent_primary'])
        self.file_path_entry.pack(side=tk.LEFT, padx=(10, 10), pady=10)
        
        # Browse button
        self.browse_btn = tk.Button(file_frame, text="Browse", command=self.browse_file,
                                   bg=self.colors['button_primary'], 
                                   fg=self.colors['text_primary'], 
                                   font=("Helvetica", 10, "bold"),
                                   relief=tk.FLAT, padx=20, cursor="hand2")
        self.browse_btn.pack(side=tk.RIGHT, padx=(0, 10), pady=10)
        
        # Convert button
        self.convert_btn = tk.Button(main_frame, text="Convert to Strapi Layout", 
                                    command=self.convert_document,
                                    bg=self.colors['button_secondary'], 
                                    fg=self.colors['text_primary'], 
                                    font=("Helvetica", 14, "bold"),
                                    relief=tk.FLAT, padx=30, pady=10, cursor="hand2")
        self.convert_btn.pack(pady=20)
        
        # Results frame
        results_frame = tk.LabelFrame(main_frame, text="Extracted Data", 
                                    bg=self.colors['bg_secondary'], 
                                    fg=self.colors['text_primary'],
                                    font=("Helvetica", 12, "bold"))
        results_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 20))
        
        # Create text widget with scrollbar
        text_frame = tk.Frame(results_frame, bg=self.colors['bg_secondary'])
        text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.results_text = tk.Text(text_frame, wrap=tk.WORD, font=("Consolas", 10),
                                   bg=self.colors['bg_accent'], 
                                   fg=self.colors['text_primary'],
                                   insertbackground=self.colors['accent_primary'],
                                   selectbackground=self.colors['accent_primary'],
                                   selectforeground=self.colors['bg_primary'])
        scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.results_text.yview)
        self.results_text.configure(yscrollcommand=scrollbar.set)
        
        self.results_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready to convert Word documents")
        self.status_bar = tk.Label(main_frame, textvariable=self.status_var, 
                                 relief=tk.FLAT, anchor=tk.W, 
                                 bg=self.colors['bg_secondary'], 
                                 fg=self.colors['text_secondary'],
                                 font=("Helvetica", 10))
        self.status_bar.pack(fill=tk.X, pady=(10, 0))
    
    def browse_file(self):
        """Open file dialog to select Word document"""
        file_path = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.file_path_var.set(file_path)
            self.status_var.set(f"Selected file: {os.path.basename(file_path)}")
            self.status_bar.configure(fg=self.colors['success'])
    
    def extract_table_data(self, doc):
        """Extract data from the table at the beginning of the document"""
        table_data = {}
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    field_name = cells[0].strip()
                    field_value = cells[1].strip()
                    
                    # Check if this field is in our template
                    for template_field in self.template_fields:
                        if template_field.lower() in field_name.lower():
                            table_data[template_field] = field_value
                            break
        
        return table_data
    
    def extract_content(self, doc):
        """Extract the content draft that follows the table"""
        content = []
        table_found = len(doc.tables) > 0  # If there are tables, we assume the template table is present
        content_section_found = False
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # If we have tables, we can start looking for content
            if table_found:
                # Look for the "Content Draft" section
                if text.lower() == 'content draft':
                    content_section_found = True
                    continue
                
                # Skip other headers and titles
                if text.lower() in ['blog template document', 'this document follows the blog template format with a table containing metadata fields followed by the content draft.']:
                    continue
                
                # If we've found the content section, start collecting content
                if content_section_found:
                    content.append(text)
        
        return '\n\n'.join(content)
    
    def generate_strapi_layout(self, table_data, content):
        """Generate Strapi layout JSON"""
        strapi_data = {
            "data": {
                "title": table_data.get("Working Title", ""),
                "author": table_data.get("Author", ""),
                "topic": table_data.get("Topic", ""),
                "blogCategory": table_data.get("Blog Category", ""),
                "targetKeywords": table_data.get("Target Keywords", ""),
                "targetAudience": table_data.get("Target Audience", ""),
                "funnelStage": table_data.get("Funnel Stage", ""),
                "cta": table_data.get("CTA", ""),
                "metaDescription": table_data.get("Working Meta Description", ""),
                "content": content,
                "publishedAt": datetime.now().isoformat(),
                "createdAt": datetime.now().isoformat(),
                "updatedAt": datetime.now().isoformat()
            }
        }
        
        return strapi_data
    
    def convert_document(self):
        """Convert the selected Word document to Strapi layout"""
        file_path = self.file_path_var.get()
        
        if not file_path:
            messagebox.showerror("Error", "Please select a Word document first.")
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror("Error", "Selected file does not exist.")
            return
        
        try:
            self.status_var.set("Processing document...")
            self.status_bar.configure(fg=self.colors['warning'])
            self.root.update()
            
            # Load the Word document
            doc = Document(file_path)
            
            # Extract table data
            table_data = self.extract_table_data(doc)
            
            # Extract content
            content = self.extract_content(doc)
            
            # Generate Strapi layout
            strapi_data = self.generate_strapi_layout(table_data, content)
            
            # Display results
            self.display_results(table_data, content, strapi_data)
            
            # Save to file
            self.save_strapi_file(strapi_data, file_path)
            
            self.status_var.set("Conversion completed successfully!")
            self.status_bar.configure(fg=self.colors['success'])
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            self.status_var.set("Error during conversion")
            self.status_bar.configure(fg=self.colors['error'])
    
    def display_results(self, table_data, content, strapi_data):
        """Display the extracted data in the results area"""
        self.results_text.delete(1.0, tk.END)
        
        # Display extracted table data
        self.results_text.insert(tk.END, "EXTRACTED TABLE DATA:\n")
        self.results_text.insert(tk.END, "=" * 50 + "\n\n")
        
        for field, value in table_data.items():
            self.results_text.insert(tk.END, f"{field}: {value}\n")
        
        self.results_text.insert(tk.END, "\n" + "=" * 50 + "\n\n")
        
        # Display content preview
        self.results_text.insert(tk.END, "CONTENT PREVIEW:\n")
        self.results_text.insert(tk.END, "=" * 50 + "\n\n")
        
        content_preview = content[:500] + "..." if len(content) > 500 else content
        self.results_text.insert(tk.END, content_preview)
        
        self.results_text.insert(tk.END, "\n\n" + "=" * 50 + "\n\n")
        
        # Display Strapi JSON
        self.results_text.insert(tk.END, "STRAPI LAYOUT JSON:\n")
        self.results_text.insert(tk.END, "=" * 50 + "\n\n")
        self.results_text.insert(tk.END, json.dumps(strapi_data, indent=2))
    
    def save_strapi_file(self, strapi_data, original_file_path):
        """Save the Strapi layout to a JSON file"""
        # Generate output filename
        base_name = os.path.splitext(os.path.basename(original_file_path))[0]
        output_dir = os.path.dirname(original_file_path)
        output_file = os.path.join(output_dir, f"{base_name}_strapi.json")
        
        # Save the file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(strapi_data, f, indent=2, ensure_ascii=False)
        
        messagebox.showinfo("Success", f"Strapi layout saved to:\n{output_file}")
    
    def run(self):
        """Run the application"""
        self.root.mainloop()


def main():
    """Main function"""
    app = WordToStrapiConverter()
    app.run()


if __name__ == "__main__":
    main() 