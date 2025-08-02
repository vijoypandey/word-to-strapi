#!/usr/bin/env python3
"""
Sample Word Document Generator
Creates a sample Word document following the blog template format
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_document():
    """Create a sample Word document with the blog template"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Blog Template Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add description
    doc.add_paragraph('This document follows the blog template format with a table containing metadata fields followed by the content draft.')
    doc.add_paragraph('')
    
    # Create the metadata table
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    
    # Define the template fields and sample values
    fields = [
        ("Working Title", "How to Build a Successful Blog in 2024"),
        ("Author", "Sarah Johnson"),
        ("Topic", "Blogging and Content Marketing"),
        ("Blog Category", "Digital Marketing"),
        ("Target Keywords", "blogging tips, content marketing, successful blog"),
        ("Target Audience", "Aspiring bloggers and content creators"),
        ("Funnel Stage", "Consideration"),
        ("CTA", "Start Your Blog Today"),
        ("Working Meta Description", "Learn the essential steps to build a successful blog in 2024. Discover proven strategies for content creation, audience building, and monetization.")
    ]
    
    # Populate the table
    for i, (field, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = field
        row.cells[1].text = value
    
    # Add some spacing
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Add content section header
    content_header = doc.add_heading('Content Draft', level=1)
    
    # Add sample content
    doc.add_paragraph('Building a successful blog in 2024 requires more than just writing good content. You need a strategic approach that combines quality content creation, audience engagement, and effective monetization strategies.')
    
    doc.add_paragraph('In this comprehensive guide, we\'ll walk you through the essential steps to create a blog that not only attracts readers but also generates sustainable income.')
    
    # Add subheading
    doc.add_heading('1. Choose Your Niche', level=2)
    doc.add_paragraph('The first step in building a successful blog is choosing the right niche. Your niche should be something you\'re passionate about and that has a viable audience. Consider factors like:')
    
    # Add bullet points
    doc.add_paragraph('• Your personal interests and expertise', style='List Bullet')
    doc.add_paragraph('• Market demand and audience size', style='List Bullet')
    doc.add_paragraph('• Competition level and monetization potential', style='List Bullet')
    doc.add_paragraph('• Long-term sustainability and growth opportunities', style='List Bullet')
    
    doc.add_heading('2. Set Up Your Blog Platform', level=2)
    doc.add_paragraph('Once you\'ve chosen your niche, it\'s time to set up your blog platform. WordPress remains the most popular choice due to its flexibility and extensive plugin ecosystem.')
    
    doc.add_paragraph('Key considerations for your platform choice include:')
    doc.add_paragraph('• Ease of use and learning curve', style='List Bullet')
    doc.add_paragraph('• Customization options and themes', style='List Bullet')
    doc.add_paragraph('• SEO capabilities and optimization tools', style='List Bullet')
    doc.add_paragraph('• Monetization features and integrations', style='List Bullet')
    
    doc.add_heading('3. Create Quality Content Consistently', level=2)
    doc.add_paragraph('Content is the backbone of any successful blog. Focus on creating high-quality, valuable content that addresses your audience\'s needs and pain points.')
    
    doc.add_paragraph('Develop a content calendar and stick to a regular publishing schedule. This helps build trust with your audience and improves your search engine rankings.')
    
    doc.add_heading('4. Build Your Audience', level=2)
    doc.add_paragraph('Growing your audience requires a multi-channel approach:')
    
    doc.add_paragraph('• Optimize your content for search engines (SEO)', style='List Bullet')
    doc.add_paragraph('• Leverage social media platforms to promote your content', style='List Bullet')
    doc.add_paragraph('• Engage with your audience through comments and social media', style='List Bullet')
    doc.add_paragraph('• Build an email list for direct communication', style='List Bullet')
    doc.add_paragraph('• Collaborate with other bloggers in your niche', style='List Bullet')
    
    doc.add_heading('5. Monetize Your Blog', level=2)
    doc.add_paragraph('There are several ways to monetize your blog:')
    
    doc.add_paragraph('• Display advertising and Google AdSense', style='List Bullet')
    doc.add_paragraph('• Affiliate marketing and product recommendations', style='List Bullet')
    doc.add_paragraph('• Sponsored content and brand partnerships', style='List Bullet')
    doc.add_paragraph('• Digital products and online courses', style='List Bullet')
    doc.add_paragraph('• Membership sites and premium content', style='List Bullet')
    
    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph('Building a successful blog takes time, effort, and consistency. Focus on providing value to your audience, and the results will follow. Remember that success doesn\'t happen overnight, but with dedication and the right strategies, you can create a thriving blog that generates both traffic and income.')
    
    doc.add_paragraph('Ready to start your blogging journey? Take the first step today and begin building the blog of your dreams!')
    
    # Save the document
    doc.save('sample_blog_template.docx')
    print("Sample Word document created: sample_blog_template.docx")

if __name__ == "__main__":
    create_sample_document() 