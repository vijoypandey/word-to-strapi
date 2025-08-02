#!/usr/bin/env python3
"""
Command-line Word to Strapi Converter
Converts Word documents to Strapi layouts via command line
"""

import argparse
import json
import os
import sys
from docx import Document
from datetime import datetime


class CLIWordToStrapiConverter:
    def __init__(self):
        # Template fields
        self.template_fields = [
            "Working Title",
            "Author", 
            "Topic",
            "Blog Category",
            "Target Keywords",
            "Target Audience",
            "Funnel Stage",
            "CTA",
            "Working Meta Description"
        ]
    
    def extract_table_data(self, doc):
        """Extract data from the table at the beginning of the document"""
        table_data = {}
        
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    field_name = cells[0].strip()
                    field_value = cells[1].strip()
                    
                    # Check if this field is in our template
                    for template_field in self.template_fields:
                        if template_field.lower() in field_name.lower():
                            table_data[template_field] = field_value
                            break
        
        return table_data
    
    def extract_content(self, doc):
        """Extract the content draft that follows the table"""
        content = []
        table_found = len(doc.tables) > 0  # If there are tables, we assume the template table is present
        content_section_found = False
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # If we have tables, we can start looking for content
            if table_found:
                # Look for the "Content Draft" section
                if text.lower() == 'content draft':
                    content_section_found = True
                    continue
                
                # Skip other headers and titles
                if text.lower() in ['blog template document', 'this document follows the blog template format with a table containing metadata fields followed by the content draft.']:
                    continue
                
                # If we've found the content section, start collecting content
                if content_section_found:
                    content.append(text)
        
        return '\n\n'.join(content)
    
    def generate_strapi_layout(self, table_data, content):
        """Generate Strapi layout JSON"""
        strapi_data = {
            "data": {
                "title": table_data.get("Working Title", ""),
                "author": table_data.get("Author", ""),
                "topic": table_data.get("Topic", ""),
                "blogCategory": table_data.get("Blog Category", ""),
                "targetKeywords": table_data.get("Target Keywords", ""),
                "targetAudience": table_data.get("Target Audience", ""),
                "funnelStage": table_data.get("Funnel Stage", ""),
                "cta": table_data.get("CTA", ""),
                "metaDescription": table_data.get("Working Meta Description", ""),
                "content": content,
                "publishedAt": datetime.now().isoformat(),
                "createdAt": datetime.now().isoformat(),
                "updatedAt": datetime.now().isoformat()
            }
        }
        
        return strapi_data
    
    def convert_document(self, input_file, output_file=None, verbose=False):
        """Convert Word document to Strapi layout"""
        
        if not os.path.exists(input_file):
            print(f"Error: File '{input_file}' does not exist.")
            return False
        
        try:
            if verbose:
                print(f"Processing document: {input_file}")
            
            # Load the Word document
            doc = Document(input_file)
            
            # Extract table data
            table_data = self.extract_table_data(doc)
            
            if verbose:
                print(f"Extracted {len(table_data)} fields from table")
                for field, value in table_data.items():
                    print(f"  {field}: {value}")
            
            # Extract content
            content = self.extract_content(doc)
            
            if verbose:
                print(f"Extracted content length: {len(content)} characters")
                print(f"Content preview: {content[:100]}...")
            
            # Generate Strapi layout
            strapi_data = self.generate_strapi_layout(table_data, content)
            
            # Determine output file
            if output_file is None:
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                output_file = f"{base_name}_strapi.json"
            
            # Save to file
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(strapi_data, f, indent=2, ensure_ascii=False)
            
            if verbose:
                print(f"Strapi layout saved to: {output_file}")
            else:
                print(f"Successfully converted '{input_file}' to '{output_file}'")
            
            return True
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return False


def main():
    """Main function"""
    parser = argparse.ArgumentParser(description='Convert Word documents to Strapi layouts')
    parser.add_argument('input_file', help='Input Word document (.docx)')
    parser.add_argument('-o', '--output', help='Output JSON file (default: input_name_strapi.json)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose output')
    
    args = parser.parse_args()
    
    converter = CLIWordToStrapiConverter()
    success = converter.convert_document(args.input_file, args.output, args.verbose)
    
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main() 